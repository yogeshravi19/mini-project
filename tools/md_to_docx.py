from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


IMG_RE = re.compile(r"!\[(?P<alt>.*?)\]\((?P<src>.*?)\)\s*$")


def md_to_docx(md_path: Path, out_path: Path) -> None:
    base_dir = md_path.parent
    # Map SVG figure links used in Markdown to PNG versions for DOCX embedding.
    # (python-docx cannot embed SVG images directly.)
    svg_to_png = {
        "figures/fig1_hallucination_taxonomy.svg": "figures/fig1_taxonomy.png",
        "figures/fig2_trueguard_overview.svg": "figures/fig2_overview.png",
        "figures/fig3_multi_signal_fusion.svg": "figures/fig3_fusion.png",
        "figures/fig4_closed_loop_mitigation.svg": "figures/fig4_mitigation.png",
        "figures/fig5_xai_trust_dimensions.svg": "figures/fig5_xai.png",
    }
    text = md_path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for raw in lines:
        line = raw.rstrip("\n")

        m = IMG_RE.fullmatch(line.strip())
        if m:
            alt = (m.group("alt") or "").strip()
            src = (m.group("src") or "").strip()
            src = svg_to_png.get(src, src)

            img_path = Path(src)
            if not img_path.is_absolute():
                img_path = (base_dir / img_path).resolve()

            p = doc.add_paragraph()
            run = p.add_run()
            if img_path.exists():
                run.add_picture(str(img_path), width=Inches(6.2))
            else:
                p.add_run(f"[Missing image: {src}]")

            if alt:
                cap = doc.add_paragraph(alt)
                for r in cap.runs:
                    r.italic = True
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        stripped = line.strip()

        # Very lightweight heading detection for this specific paper.
        if re.match(r"^[IVX]+\.\s+", stripped):
            doc.add_heading(stripped, level=1)
            continue
        if re.match(r"^[A-Z]\.\s+", stripped):
            doc.add_heading(stripped, level=2)
            continue

        doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == "__main__":
    md = Path(r"E:/mini-project/Toward Trustworthy Large Language Models.md")
    out = Path(r"E:/mini-project/Toward Trustworthy Large Language Models_FINAL.docx")
    md_to_docx(md, out)
    print(f"Wrote: {out}")

